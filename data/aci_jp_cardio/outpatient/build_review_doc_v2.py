#!/usr/bin/env python3
"""
v2 review request document — incorporates all 33 doctor's comments [a]-[bg]
from 2026-05-01 review round.

Adds:
- 冒頭に「前回コメントどう反映したか」サマリ表
- 各症例冒頭に「v2 で具体的に何を直したか」のミニ要約
- 各症例末尾の「先生コメント欄」(再レビュー用、空白)

Usage:
    python3.11 build_review_doc_v2.py
"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).parent
CASES_DIR = ROOT / "cases"
OUT_DOCX = ROOT / "review_request_for_specialist_v2.docx"
OUT_TXT = ROOT / "review_request_for_specialist_v2.txt"

# 8 cases in the same priority order as v1, plus per-case "v2 で直した点"
SELECTION = [
    {
        "id": "JC-EA-B",
        "priority_note": "【最優先】前回 [b]〜[t] (19 件) のコメントを反映した v2。Pattern B (口語対話) の自然化と Plan の現実化に注力。",
        "v2_changes": [
            "[b] 血液型 A → null (初診で血液型は持っていないのが普通)",
            "[c][m] 副主訴の「心拍速感」→「動悸」",
            "[d][h] 「元喫煙者」→「過去喫煙 (Brinkman 500)」",
            "[i] 社会歴のデスクワークとジョギングを別文に分割",
            "[j] 「持続的に出現」→「持続」",
            "[n] 「ジョギング中増悪なし」→「ジョギング中の増悪なし」",
            '[o] 「古典的労作性狭心症」の「古典的」を削除',
            "[p] Plan の「即時 12 誘導 ECG」→「12 誘導 ECG を記録」",
            "[q] 「トロポニン I (0/3/6 時間連続)」→「1 回採血、陽性なら数時間後追加」",
            "[r] 「SpO2 経時モニター」を Plan から削除",
            "[s] 「ACS 強疑時」→「ACS を強く疑う場合」",
            "[t] 喫煙再開しないよう指導継続を削除 (1 年弱の禁煙者には不要)",
            "[e][f][g] Pattern B の「簡単な聴診」「12 誘導心電図」「採血もお願いします」を「胸の音を聞かせてもらいます」「心電図を取らせてもらいます」「採血もさせていただきます」に書き換え",
        ],
    },
    {
        "id": "JC-EA-A",
        "priority_note": "【最優先】前回 [u]〜[ag] (7 件) のコメントを反映した v2。再診カルテの慣習を反映、検査オーダーを現実的に。",
        "v2_changes": [
            "[u] 「前回受診後より NTG ゼロ」→「PCI 後ニトログリセリン舌下投与は使用していない」",
            "[v][w] O セクションの「活気良好、苦悶様顔貌なし」を削除 (普段の再診で書かない)",
            "[x][aa][ab] O の「橈骨動脈・足背動脈触知良好」「PCI 部位異常なし」を削除 (再診で触診しない)",
            "[ac] BNP 32 の「(低値)」→「(基準内)」(医学的に正常範囲)",
            "[ad] β遮断薬・ACE 阻害薬を「心保護目的継続」と書きすぎず「慣行的継続」に (心機能良好なら必須ではない)",
            "[ae] SGLT2-i 「次回検討」→「HbA1c が今後上昇傾向となれば追加検討」(改善傾向では不要)",
            "[af] 「家庭でのコンディショニング維持」→「自宅でも運動習慣は維持されている」",
            "[ag] 「3 ヶ月毎心エコー」を削除し「症状再燃時または 12 ヶ月後」に変更",
        ],
    },
    {
        "id": "JC-AMI-B",
        "priority_note": "前回 [ah]〜[ap] (6 件) のコメントを反映した v2。家族歴の定義整理 + 検査オーダー現実化。",
        "v2_changes": [
            "[ah][aj] 家族歴から「夫: 脳梗塞」を削除 (配偶者は血縁ではない、社会歴へ移動)",
            "[ai] 既往の聞き方を「他に何か病気をされたことはありますか?」と自然な日本語に",
            "[ak] HR 56 の「徐脈」→「やや徐脈」",
            "[al] SpO2 95% の「軽度低下」を削除 (95% は正常範囲)",
            "[am][an] V3R/V4R 右側胸部誘導の「必須」→「必要時に検討」(普通の心電図で右室梗塞所見なくても疑わしい場合のみ)",
            "[ao] ECG で明らかな ST 変化があれば「採血を待たず搬送」を P に追加",
            "[ap] N-アセチルシステインを造影剤腎症予防の文脈から削除 (現状使われない)",
        ],
    },
    {
        "id": "JC-CA-B",
        "priority_note": "前回 [aq]〜[at] (4 件) のコメントを反映した v2。Pattern B の繊細な聞き方 + 降圧薬の現実的な調整。",
        "v2_changes": [
            "[aq] 母の死因確認の言い方「お母様が 70 歳代で心臓の病気で亡くなられたんですね」と相手の言葉を受け止める形に",
            "[ar][as] 「それは何という病気だったか、お時間少しいただいてもよろしいでしょうか?」と丁寧な聞き方を Pattern B に追加",
            "[at] 本日 BP 108/68 + 立ちくらみあり → テルミサルタン (ARB) を「減量 (5→2.5mg) もしくは中止指示」を Plan に明記",
        ],
    },
    {
        "id": "JC-AHF-B",
        "priority_note": "前回 [au]〜[ay] (5 件) のコメントを反映した v2。主訴の自然化 + Plan の冗長性削減。",
        "v2_changes": [
            "[au] 主訴「急に息ができない」→「急性発症の呼吸苦」",
            "[av] [医師] 「こんにちは、診察します」→「こんにちは、お話をお聞かせください」",
            "[aw] 肺炎合併では「痰の喀出ができず悪化リスク」を A/P に追加",
            "[ax] **Plan を 10 → 6 項目に削減**、Assessment と重複する病態解説を P から除去",
            "[ay] トロポニン I 「0/3/6 時間連続」→「1 回 + 必要時」(ACS でもタコツボでも上昇するため連続フォロー不要)",
        ],
    },
    {
        "id": "JC-AS-B",
        "priority_note": "前回 [az]〜[bf] (7 件) のコメントを反映した v2。専門用語を患者・家族の口語から外し、Plan を圧縮。",
        "v2_changes": [
            "[az] 家族歴「父: 不明 (戦死)」を削除 (医学的に意味がない)",
            "[ba] 「喫煙: なし (生涯非喫煙)」→「喫煙歴: なし」",
            "[bb] 長女の発言「MMSE 23 点で軽度認知症」「要介護 1 で」→「物忘れの検査で軽度の認知症と言われた」「介護保険を使わせてもらってる」(専門スコアを家族が口にしない)",
            "[bc] [医師] の問診段階での詳細病態説明 (「大動脈弁狭窄症という弁の病気を強く疑っています、突然死のリスクも...」) を「これから検査をしていきます。検査結果を見て、またご説明させていただきますね」に",
            "[bd] 主訴「失神・労作時息切れ」→「たまに気を失う」「坂道で息が切れる」(専門的すぎる病名を主訴に書かない)",
            "[be] 受付バイタル「臥位 BP 108/68 (やや低め)」→「やや低め」表現を削除 (正常範囲)",
            "[bf] **Plan を圧縮**、Assessment との重複を削除",
        ],
    },
    {
        "id": "JC-VSA-B",
        "priority_note": "前回 [bg] (1 件) のコメントを反映した v2 + 共通テーマ反映。",
        "v2_changes": [
            "[bg] [医師] 「早死にした親族とかは?」→「若くして亡くなられた親族の方はいらっしゃいますか?」",
            "+ 共通テーマ: blood_type→null、過去喫煙表記、Pattern B の自然化、Plan 圧縮、配偶者を家族歴から外す、SpO2 経時モニターを Plan から削除",
        ],
    },
    {
        "id": "JC-SCP-B",
        "priority_note": "前回コメント未受領 (時間切れ)。LESSONS_LEARNED.md の共通テーマを全面反映した v2。",
        "v2_changes": [
            "blood_type → null",
            "Pattern B の患者向け会話自然化 (12 誘導/簡単な聴診/お願いします/診察します を回避)",
            "Plan を 11 → 7 項目に削減、A との重複削除",
            "配偶者を家族歴から外す",
            "is_negative_control: true 維持、循環器薬は処方しない方針維持",
        ],
    },
]


# 共通テーマ集計 (冒頭サマリ用)
COMMON_THEMES = [
    ("カルテ表現の慣習", "「即時」「軽度低下」「やや低め」「活気」「全身状態」「苦悶様顔貌なし」「強疑時」を削除/書き換え。「持続的に出現」→「持続」、「家庭でのコンディショニング維持」→「自宅でも運動習慣は維持されている」、「元喫煙者」→「過去喫煙」、「生涯非喫煙」→「なし」", "[p][al][be][v][w][s][j][af][d][h][ba][o]"),
    ("患者向け会話の自然化", "「12 誘導心電図」「簡単な聴診」「採血もお願いします」「診察します」「早死にした親族」を回避し、「胸の音を聞かせてもらいます」「採血もさせていただきます」「お話をお聞かせください」「若くして亡くなられた親族の方」に。問診段階での詳細病態説明 (「強く疑っています、突然死のリスクも」等) は「検査結果を見てご説明します」に", "[e][f][g][av][bg][bc]"),
    ("Plan 構造の現実化", "全 22 例で Plan を 5-9 項目に圧縮 (mean 7.5)、A セクションとの重複病態解説を P から削除。「Plan は Plan だけを書く」「日常でここまでカルテは書けない」原則の反映", "[ax][bf]"),
    ("家族歴の定義整理", "配偶者は家族歴ではない (血縁ではない) → 社会歴に移動。「父: 戦死」など医学的に意味のない情報を削除", "[ah][aj][az]"),
    ("検査オーダーの現実度", "atypical 胸痛で「トロポニン 0/3/6 時間連続」を「1 回 + 陽性時追加」に変更。SpO2 経時モニター・3 ヶ月毎心エコー・再診での橈骨動脈/足背動脈触診・V3R/V4R 右側胸部誘導のルーチン化・N-アセチルシステインによる造影剤腎症予防 を全削除。ECG で明らかな ST 変化があれば採血を待たず搬送", "[q][r][y][ag][x][aa][ab][am][an][ap][ao]"),
    ("患者プロファイル", "blood_type は初診で持っていないのが普通 → 全例 null。社会歴で職業/運動/喫煙/飲酒を別文に分割", "[b][i]"),
    ("既往の聞き方", "「他に何か病気をされたことはありますか?」と自然な日本語に", "[ai]"),
    ("数値の解釈", "BNP 32 は基準内 (システム表示は高値マーク)、HR 56 は「やや徐脈」、SpO2 95% は正常範囲なので「軽度低下」と書かない", "[ac][ak][al]"),
    ("薬剤調整", "心機能良好なら β遮断薬・ACE 阻害薬は必須ではない (慣行的継続は OK)、HbA1c 改善傾向で SGLT2-i 追加は不要、ARB 内服中で立ちくらみあれば減量 5→2.5mg or 中止指示", "[ad][ae][at]"),
    ("VSA 治療", "VSA で β遮断薬は禁忌 (冠攣縮悪化のため)、Ca 拮抗薬 + 硝酸薬第一選択 を A/P で明示", "(LESSONS_LEARNED.md 4 項)"),
    ("細部修正", "「ジョギング中増悪なし」→「ジョギング中の増悪なし」、「古典的労作性狭心症」の「古典的」削除、「禁煙 1 年弱の患者に強い禁煙指導」削除、肺炎合併で「痰の喀出できず悪化リスク」を A/P に明記、母の死因確認は丁寧に", "[n][o][t][aw][aq][ar][as]"),
]


def load_cases() -> dict:
    cases = {}
    for f in CASES_DIR.glob("JC-*.json"):
        with open(f, encoding="utf-8") as fp:
            c = json.load(fp)
        cases[c["encounter_id"]] = c
    return cases


def case_to_text_blocks(c: dict) -> list:
    """Return list of (style, text) tuples for a case."""
    blocks = []
    p = c.get("patient", {})
    e = c.get("encounter", {})
    rv = c.get("reception_vitals", {})
    kf = c.get("key_facts", {})

    scenario_label = "シナリオ B (初診 walk-in)" if c["scenario"] == "shoshin_walkin" else "シナリオ A (再診/紹介状)"
    diff_label = "陰性対照" if c.get("is_negative_control") else f"難度: {c['difficulty']}"
    blocks.append(("h2", f"{c['encounter_id']} — {c['disease_label_jp']}"))
    blocks.append(("subtitle", f"{scenario_label} / {diff_label}"))

    # Patient
    blocks.append(("h3", "■ 患者情報"))
    pt_lines = [
        f"・年齢/性別: {p.get('age')} 歳 {p.get('gender', '')}" + (f" (血液型 {p.get('blood_type')})" if p.get("blood_type") else ""),
        f"・主訴: {e.get('chief_complaint', '')}",
    ]
    if e.get("secondary_complaints"):
        pt_lines.append(f"・副主訴: {'、'.join(e['secondary_complaints'])}")
    pt_lines += [
        f"・既往: {'、'.join(p.get('comorbidities', [])) or '—'}",
        f"・持参薬: {'、'.join(p.get('current_medications', [])) or '—'}",
        f"・アレルギー: {'、'.join(p.get('allergies', [])) or '—'}",
        f"・家族歴: {'、'.join(p.get('family_history', [])) or '—'}",
        f"・社会歴: {p.get('social_history', '—')}",
        f"・受診: {e.get('encounter_date', '')} / {e.get('department', '')} / {e.get('type', '')}",
    ]
    for line in pt_lines:
        blocks.append(("body", line))

    # Reception vitals
    if rv:
        blocks.append(("h3", "■ 受付バイタル"))
        v_parts = []
        if rv.get("BP_sys") is not None:
            v_parts.append(f"BP {rv['BP_sys']}/{rv['BP_dia']} mmHg")
        if rv.get("HR") is not None:
            v_parts.append(f"HR {rv['HR']}/min")
        if rv.get("SpO2") is not None:
            v_parts.append(f"SpO2 {rv['SpO2']}%")
        if rv.get("RR") is not None:
            v_parts.append(f"RR {rv['RR']}")
        if rv.get("BT") is not None:
            v_parts.append(f"BT {rv['BT']}℃")
        blocks.append(("body", "・" + " / ".join(v_parts)))

    # Input
    blocks.append(("h3", "■ 入力 (SLM への問診)"))
    if c["scenario"] == "shoshin_walkin":
        blocks.append(("body_label", "[口語対話 Pattern B]"))
        blocks.append(("body_quoted", c.get("input_pattern_B", "")))
    else:
        ipa = c.get("input_pattern_A", {}) or {}
        ps = ipa.get("physician_summary", {}) or {}
        for label, key in [
            ("医師サマリ (現病歴・既往整理)", "raw_text"),
            ("お薬手帳", "medication_list"),
            ("診察所見", "exam_findings"),
            ("検査結果", "lab_results"),
        ]:
            v = ps.get(key)
            if v:
                blocks.append(("body_label", f"[{label}]"))
                blocks.append(("body_quoted", v))

    # Reference SOAP
    blocks.append(("h3", "■ 正解 (reference) SOAP"))
    soap = c.get("reference_soap", {})
    for sec_label, sec in [("S (Subjective)", "S"), ("O (Objective)", "O"),
                            ("A (Assessment)", "A"), ("P (Plan)", "P")]:
        if soap.get(sec):
            blocks.append(("body_label", f"[{sec_label}]"))
            blocks.append(("body_quoted", soap[sec]))

    # Admission summary
    if c.get("reference_admission_summary"):
        blocks.append(("h3", "■ 入院時サマリ"))
        blocks.append(("body_quoted", c["reference_admission_summary"]))

    # Key facts (compact)
    blocks.append(("h3", "■ key_facts (eval runner が機械採点する正解集合)"))
    kf_lines = []
    if kf.get("expected_triage"):
        kf_lines.append(f"・期待 triage 判断: {kf['expected_triage']}")
    if kf.get("must_not_miss"):
        kf_lines.append(f"・Must-not-miss: {'、'.join(kf['must_not_miss'])}")
    if kf.get("differential_diagnoses"):
        kf_lines.append(f"・鑑別診断: {'、'.join(kf['differential_diagnoses'])}")
    if kf.get("appropriate_next_tests"):
        kf_lines.append(f"・推奨検査: {'、'.join(kf['appropriate_next_tests'])}")
    if kf.get("diagnoses"):
        kf_lines.append(f"・確定診断: {'、'.join(kf['diagnoses'])}")
    if kf.get("medications_to_start"):
        kf_lines.append(f"・開始薬: {'、'.join(kf['medications_to_start'])}")
    if kf.get("medications_to_continue"):
        kf_lines.append(f"・継続薬: {'、'.join(kf['medications_to_continue'])}")
    if kf.get("medications_to_stop"):
        kf_lines.append(f"・中止薬: {'、'.join(kf['medications_to_stop'])}")
    if kf.get("disposition"):
        kf_lines.append(f"・転帰: {kf['disposition']}")
    for line in kf_lines:
        blocks.append(("body", line))

    return blocks


def write_docx(cases: dict, selection: list):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10.5)

    doc.add_heading("ACI-JP-Cardio 外来カルテ生成ベンチマーク — 専門医レビュー (v2 修正版)", level=0)

    # Intro
    doc.add_paragraph(
        "前回 (2026-04-29 v1) いただいた 33 件のコメント [a] 〜 [bg] を全例に反映した v2 を作成しました。"
        "ご多忙のところ恐れ入りますが、修正が想定通りに反映されているかご確認いただけますと幸いです。"
    )
    doc.add_paragraph(
        "前回と同じ 8 例を抜粋しています。"
        "各症例の冒頭に「v2 で具体的に何を直したか」を太字で列挙していますので、"
        "そこを見て「OK」または「まだ気になる箇所」をお知らせください。"
    )

    # === なぜレビューが必要か (評価指標の説明) ===
    doc.add_heading("なぜ専門医レビューが必要か (評価指標の説明)", level=1)
    doc.add_paragraph(
        "ご記入いただく reference SOAP は、SLM (Small Language Model) を採点するときの "
        "「物差し」になります。物差しが歪んでいるとすべての評価が信頼できなくなるため、"
        "医学的に妥当な reference を作っていただくことが本ベンチマークの土台です。"
    )

    doc.add_heading("採点の流れ (1 症例あたり)", level=2)
    p = doc.add_paragraph()
    p.add_run("【入力】問診テキスト (本書の各症例 input_pattern_A/B)\n").bold = True
    p.add_run("    ↓\n")
    p.add_run("【SLM が出力】Qwen3.5 ベースの SLM が SOAP を生成\n").bold = True
    p.add_run("    ↓\n")
    p.add_run("【採点】SLM 出力 と reference SOAP を 6-7 種類の指標で機械比較\n").bold = True
    p.add_run("    ↓\n")
    p.add_run("【最終スコア】各指標を平均した composite スコアが「このモデルの実力」\n").bold = True

    doc.add_heading("採点指標の一覧 (噛み砕いた説明)", level=2)

    metrics_table = doc.add_table(rows=1, cols=3)
    metrics_table.style = "Light Grid Accent 1"
    hdr = metrics_table.rows[0].cells
    hdr[0].text = "指標"
    hdr[1].text = "何を測るか"
    hdr[2].text = "ハック耐性 / 弱点"

    metrics_rows = [
        ("ROUGE-L (単語一致)",
         "AI 出力と reference の「同じ単語が並ぶ最長部分列」の一致率。"
         "例: reference「アスピリン 100mg を投与」と AI「アスピリン 100mg を経口投与」"
         "→ 高一致。MeCab で分かち書きしてから計算。",
         "言い換え (動悸 ↔ 胸がドキドキ) に弱く、同じ意味でも別単語なら 0 になる"),
        ("BERTScore (意味類似度)",
         "BERT という日本語 AI に各単語を埋め込み (embedding) → 意味的に近い単語マッチを採点。"
         "「動悸」と「胸がドキドキする」を意味的に近いと判定できる。",
         "BERT 自体の医学知識が浅いと細かい医学的差異 (Killip I vs II など) が見えない"),
        ("薬剤 F1 (我々の novel 指標)",
         "reference の key_facts.medications_to_start/continue にある薬剤と、AI が出力に含めた薬剤を比較。"
         "同義語辞書 (例: アムロジピン = ノルバスク = amlodipine) で正規化。"
         "AI が架空の薬名を出力すると即マイナス点。",
         "辞書に登録されていないマイナーな薬は評価できない"),
        ("診断 F1",
         "reference の key_facts.diagnoses に含まれる病名 (例: 急性前壁心筋梗塞 / Killip II) と"
         "AI 出力に含まれる病名を比較。日英表記揺れ (STEMI = ST 上昇型心筋梗塞) を吸収。",
         "stage 分類など細かい修飾語の評価は粗い"),
        ("バイタル一致率",
         "reference の vitals/labs にある数値 (BP 152/94, HR 98 等) を AI が出力に含めているかを"
         "±10% 許容で一致判定。",
         "絶対値ベースなので「軽度」「中等度」のような言語的記述は評価できない"),
        ("Opus-as-judge (5 軸 rubric)",
         "Claude Opus に reference + AI 出力を渡して、医学的妥当性 / 情報網羅性 / 日本語自然さ / "
         "ハルシネーション無さ / フォーマット遵守 の 5 軸で 1-5 点採点。"
         "上記の機械指標で見えない「総合的な質感」を評価。",
         "Opus 自体の医学知識依存。確定診断系は精度高いが、薬剤量微調整等は判断ぶれあり"),
    ]
    for name, what, weak in metrics_rows:
        row = metrics_table.add_row().cells
        row[0].text = name
        row[1].text = what
        row[2].text = weak

    doc.add_heading("具体例 — 4B SOAP が JC-AMI-S を解いた実測値", level=2)
    p = doc.add_paragraph()
    p.add_run("AI が生成した SOAP に対して、reference (先生に作っていただいた正解) と比較した結果:\n").italic = True

    example_table = doc.add_table(rows=1, cols=3)
    example_table.style = "Light Grid Accent 1"
    hdr = example_table.rows[0].cells
    hdr[0].text = "指標"
    hdr[1].text = "値"
    hdr[2].text = "解釈"
    example_rows = [
        ("ROUGE-L", "0.46", "単語ベースで約半分一致 (悪くない)"),
        ("BERTScore F1", "0.81", "意味的にはかなり一致"),
        ("薬剤 F1", "0.66", "アスピリン / クロピドグレル / ヘパリン は当てたが、ビソプロロールを漏らす"),
        ("診断 F1", "0.50", "「急性前壁心筋梗塞」は当たり、「Killip II」を漏らす"),
        ("バイタル一致率", "67%", "BP / HR / SpO2 は記載あり、トロポニン / CK-MB を漏らす"),
        ("Opus judge avg", "4.0/5", "medical=4, completeness=3, naturalness=5, hallucination=4, format=5"),
        ("Composite", "0.61", "「この症例での 4B SOAP の総合点」"),
    ]
    for name, val, interp in example_rows:
        row = example_table.add_row().cells
        row[0].text = name
        row[1].text = val
        row[2].text = interp

    doc.add_paragraph(
        "22 例分の Composite を平均すると「そのモデル構成の総合スコア」になります。"
        "現状 4B モデル: 0.389 / 9B モデル: 0.415 と、9B のほうが約 7% 高い結果が出ています。"
    )

    doc.add_heading("先生のレビューがどう効くか", level=2)
    p = doc.add_paragraph()
    p.add_run("もし reference の S/O/A/P や key_facts に医学的に誤った内容があると…\n").bold = True
    bullet_examples = [
        "AI が「禁忌のβ遮断薬を VSA に処方」と出力 → reference にも誤って同じ処方が書かれていれば、薬剤 F1 で「正解扱い」されて満点になり、本来検出すべき危険な誤りがベンチマークから見えなくなる",
        "AI が「Killip III と判定」と出力 → reference の Killip 分類が間違っていれば、AI の誤りが「正解扱い」される",
        "AI が「実在しない架空薬」を出力 → reference に書かれていなければ薬剤 F1 で false positive 扱いされ正しくマイナス評価される",
    ]
    for ex in bullet_examples:
        doc.add_paragraph(ex, style="List Bullet")
    doc.add_paragraph(
        "つまり「reference の質 = ベンチマーク全体の信頼性」ということで、"
        "今回の専門医レビューは AI を測るより前に「物差し自体を校正する」作業になります。"
    )

    # Common themes summary
    doc.add_heading("共通テーマの修正サマリ (前回 33 件のコメントを 11 テーマに集約)", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "テーマ"
    hdr[1].text = "反映内容"
    hdr[2].text = "該当コメント"
    for h, body, cmts in COMMON_THEMES:
        row = table.add_row().cells
        row[0].text = h
        row[1].text = body
        row[2].text = cmts

    # Number summary
    doc.add_paragraph(
        "上記テーマは 22 例全例 (= 8 例の倍以上) に反映済です。"
        "本書では特に修正密度の高い 8 例を抜粋して掲載します。"
    )

    # How to provide feedback
    doc.add_heading("修正コメントの書き方 (前回と同じ)", level=1)
    doc.add_paragraph(
        "問題なければ「OK」、まだ気になる箇所があれば下の「先生コメント欄」に "
        "Google Docs のコメント機能で記載してください。"
    )

    # Selection summary
    doc.add_heading("レビュー対象 8 件 (前回と同じ)", level=1)
    summary_para = doc.add_paragraph()
    for i, sel in enumerate(selection, 1):
        c = cases.get(sel["id"])
        if not c:
            continue
        summary_para.add_run(f"{i}. {sel['id']} — {c['disease_label_jp']}\n").bold = True

    # Per-case content
    for i, sel in enumerate(selection, 1):
        c = cases.get(sel["id"])
        if not c:
            continue
        doc.add_page_break()

        doc.add_heading(f"症例 {i} / 8: {sel['id']}", level=1)

        # Priority note
        p = doc.add_paragraph()
        run = p.add_run(sel["priority_note"])
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        # v2 changes — bulleted, in green
        doc.add_heading("v2 で具体的に直した点", level=3)
        for change in sel["v2_changes"]:
            p = doc.add_paragraph(change, style="List Bullet")
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x00, 0x60, 0x40)

        # Case content
        doc.add_heading("v2 内容 (修正反映後)", level=3)
        for style_, text in case_to_text_blocks(c):
            if style_ == "h2":
                doc.add_heading(text, level=2)
            elif style_ == "h3":
                doc.add_heading(text, level=4)
            elif style_ == "subtitle":
                p = doc.add_paragraph()
                p.add_run(text).italic = True
            elif style_ == "body":
                doc.add_paragraph(text)
            elif style_ == "body_label":
                p = doc.add_paragraph()
                p.add_run(text).bold = True
            elif style_ == "body_quoted":
                p = doc.add_paragraph(text)
                p.paragraph_format.left_indent = Cm(0.5)

        # Reviewer feedback area
        doc.add_heading("先生コメント欄 (再レビュー)", level=3)
        doc.add_paragraph("(v2 で直っていれば「OK」とご記載ください。まだ気になる箇所があれば具体的にご指摘いただけますと、再修正いたします)")
        doc.add_paragraph(" ")
        doc.add_paragraph(" ")
        doc.add_paragraph(" ")

    doc.save(OUT_DOCX)
    print(f"wrote {OUT_DOCX} ({OUT_DOCX.stat().st_size / 1024:.1f} KB)")


def write_txt(cases: dict, selection: list):
    """Plain text fallback."""
    lines = []
    sep = "=" * 70
    sep_l = "-" * 70

    lines.append(sep)
    lines.append("ACI-JP-Cardio 外来カルテ生成ベンチマーク — 専門医レビュー (v2 修正版)")
    lines.append(sep)
    lines.append("")
    lines.append("前回 (2026-04-29 v1) いただいた 33 件のコメント [a] 〜 [bg] を全例に反映した")
    lines.append("v2 を作成しました。ご多忙のところ恐れ入りますが、修正が想定通りに反映")
    lines.append("されているかご確認いただけますと幸いです。")
    lines.append("")
    lines.append("前回と同じ 8 例を抜粋しています。各症例の冒頭に「v2 で具体的に何を直したか」")
    lines.append("を列挙していますので、そこを見て「OK」または「まだ気になる箇所」をお知らせ")
    lines.append("ください。")
    lines.append("")

    # === なぜレビューが必要か ===
    lines.append(sep_l)
    lines.append("なぜ専門医レビューが必要か (評価指標の説明)")
    lines.append(sep_l)
    lines.append("")
    lines.append("ご記入いただく reference SOAP は、SLM (Small Language Model) を採点するときの")
    lines.append("「物差し」になります。物差しが歪んでいるとすべての評価が信頼できなくなるため、")
    lines.append("医学的に妥当な reference を作っていただくことが本ベンチマークの土台です。")
    lines.append("")
    lines.append("【採点の流れ (1 症例あたり)】")
    lines.append("")
    lines.append("  入力 (問診テキスト)")
    lines.append("    ↓")
    lines.append("  SLM (Qwen3.5 ベース) が SOAP を生成")
    lines.append("    ↓")
    lines.append("  SLM 出力 と reference SOAP を 6-7 種類の指標で機械比較")
    lines.append("    ↓")
    lines.append("  各指標を平均した composite スコアが「このモデルの実力」")
    lines.append("")
    lines.append("【採点指標の一覧 (噛み砕いた説明)】")
    lines.append("")
    metrics_simple = [
        ("ROUGE-L (単語一致)",
         "AI 出力と reference の同じ単語の並びの一致率を採点。MeCab で分かち書き。",
         "言い換え (動悸 ↔ 胸がドキドキ) に弱い"),
        ("BERTScore (意味類似度)",
         "日本語 BERT で単語を埋め込み (embedding) → 意味的に近い単語マッチを採点。",
         "細かい医学的差異 (Killip I vs II) は見えにくい"),
        ("薬剤 F1",
         "reference の key_facts.medications にある薬と AI 出力の薬を比較。"
         "同義語辞書 (アムロジピン = ノルバスク 等) で正規化。架空薬は即マイナス。",
         "辞書未登録の薬は評価不能"),
        ("診断 F1",
         "reference の key_facts.diagnoses にある病名と AI 出力の病名を比較。"
         "日英表記揺れ (STEMI = ST 上昇型心筋梗塞) を吸収。",
         "stage 分類など細かい修飾語の評価は粗い"),
        ("バイタル一致率",
         "reference の vitals/labs の数値を AI が出力に含めているかを ±10% 許容で判定。",
         "「軽度」「中等度」のような言語的記述は評価できない"),
        ("Opus-as-judge (5 軸 rubric)",
         "Claude Opus に reference + AI 出力を渡して 医学的妥当性 / 情報網羅性 / "
         "日本語自然さ / ハルシネーション無さ / フォーマット遵守 を 5 軸 1-5 点採点。",
         "Opus 自体の医学知識依存"),
    ]
    for name, what, weak in metrics_simple:
        lines.append(f"  ・{name}")
        lines.append(f"      何を測るか: {what}")
        lines.append(f"      弱点      : {weak}")
        lines.append("")

    lines.append("【具体例 — 4B SOAP が JC-AMI-S を解いた実測値】")
    lines.append("")
    lines.append("  ROUGE-L            = 0.46  (単語ベースで約半分一致)")
    lines.append("  BERTScore F1       = 0.81  (意味的にはかなり一致)")
    lines.append("  薬剤 F1            = 0.66  (アスピリン/クロピドグレル/ヘパリン 当て、ビソプロロール漏らす)")
    lines.append("  診断 F1            = 0.50  (急性前壁心筋梗塞 当たり、Killip II 漏らす)")
    lines.append("  バイタル一致率     = 67%   (BP/HR/SpO2 OK、トロポニン/CK-MB 漏らす)")
    lines.append("  Opus judge avg     = 4.0/5 (medical=4, completeness=3, naturalness=5, hallucination=4, format=5)")
    lines.append("  Composite          = 0.61  (この症例での 4B SOAP の総合点)")
    lines.append("")
    lines.append("22 例分の Composite を平均すると「そのモデル構成の総合スコア」になります。")
    lines.append("現状 4B モデル: 0.389 / 9B モデル: 0.415 と、9B のほうが約 7% 高い結果が出ています。")
    lines.append("")
    lines.append("【先生のレビューがどう効くか】")
    lines.append("")
    lines.append("もし reference の S/O/A/P や key_facts に医学的に誤った内容があると…")
    lines.append("")
    lines.append("  ・AI が「禁忌のβ遮断薬を VSA に処方」と出力 → reference にも誤って同じ")
    lines.append("    処方が書かれていれば薬剤 F1 で「正解扱い」されて満点になり、本来検出")
    lines.append("    すべき危険な誤りがベンチマークから見えなくなる")
    lines.append("")
    lines.append("  ・AI が「Killip III と判定」と出力 → reference の Killip 分類が間違って")
    lines.append("    いれば AI の誤りが「正解扱い」される")
    lines.append("")
    lines.append("  ・AI が「実在しない架空薬」を出力 → reference に書かれていなければ薬剤 F1")
    lines.append("    で false positive 扱いされ正しくマイナス評価される")
    lines.append("")
    lines.append("つまり「reference の質 = ベンチマーク全体の信頼性」。今回の専門医レビューは")
    lines.append("AI を測るより前に「物差し自体を校正する」作業になります。")
    lines.append("")

    lines.append(sep_l)
    lines.append("共通テーマの修正サマリ (前回 33 件のコメントを 11 テーマに集約)")
    lines.append(sep_l)
    lines.append("")
    for h, body, cmts in COMMON_THEMES:
        lines.append(f"【{h}】 (該当コメント: {cmts})")
        lines.append(body)
        lines.append("")

    lines.append("上記テーマは 22 例全例 (= 8 例の倍以上) に反映済です。")
    lines.append("本書では特に修正密度の高い 8 例を抜粋して掲載します。")
    lines.append("")

    lines.append(sep_l)
    lines.append("修正コメントの書き方 (前回と同じ)")
    lines.append(sep_l)
    lines.append("")
    lines.append("問題なければ「OK」、まだ気になる箇所があれば下の「先生コメント欄」に")
    lines.append("Google Docs のコメント機能で記載してください。")
    lines.append("")

    lines.append(sep_l)
    lines.append("レビュー対象 8 件 (前回と同じ)")
    lines.append(sep_l)
    for i, sel in enumerate(selection, 1):
        c = cases.get(sel["id"])
        if not c:
            continue
        lines.append(f"  {i}. {sel['id']} — {c['disease_label_jp']}")
    lines.append("")

    for i, sel in enumerate(selection, 1):
        c = cases.get(sel["id"])
        if not c:
            continue
        lines.append("")
        lines.append(sep)
        lines.append(f"症例 {i} / 8: {sel['id']}")
        lines.append(sep)
        lines.append("")
        lines.append("【優先度メモ】")
        lines.append(sel["priority_note"])
        lines.append("")
        lines.append("【v2 で具体的に直した点】")
        for change in sel["v2_changes"]:
            lines.append(f"  ・{change}")
        lines.append("")
        lines.append("【v2 内容 (修正反映後)】")

        for style_, text in case_to_text_blocks(c):
            if style_ == "h2":
                lines.append("")
                lines.append(text)
            elif style_ == "h3":
                lines.append("")
                lines.append(text)
            elif style_ == "subtitle":
                lines.append(f"({text})")
            elif style_ == "body":
                lines.append(text)
            elif style_ == "body_label":
                lines.append("")
                lines.append(text)
            elif style_ == "body_quoted":
                for ln in text.split("\n"):
                    lines.append("  " + ln)

        lines.append("")
        lines.append("【先生コメント欄 (再レビュー)】")
        lines.append("(v2 で直っていれば「OK」とご記載ください)")
        lines.append("")
        lines.append("")
        lines.append("")
        lines.append("")

    OUT_TXT.write_text("\n".join(lines), encoding="utf-8")
    print(f"wrote {OUT_TXT} ({OUT_TXT.stat().st_size / 1024:.1f} KB)")


def main():
    cases = load_cases()
    write_docx(cases, SELECTION)
    write_txt(cases, SELECTION)


if __name__ == "__main__":
    main()
