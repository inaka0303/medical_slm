#!/usr/bin/env python3
"""
Build a Word document for specialist physician review.

Selects 8 highest-priority cases and packages them with review guidance,
schema explanation, and case-by-case content suitable for copy-paste into
Google Docs.

Usage:
    python3.11 build_review_doc.py
"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent
CASES_DIR = ROOT / "cases"
OUT_DOCX = ROOT / "review_request_for_specialist.docx"
OUT_TXT = ROOT / "review_request_for_specialist.txt"

# 8 selected cases in priority order
SELECTION = [
    {
        "id": "JC-EA-B",
        "priority_note": "【最優先】先生から提供いただいた実問診データ (2026-03-23 胸痛例) を基に Claude がハンドメイド。スキーマ検証 + 内容妥当性をご確認ください。",
    },
    {
        "id": "JC-EA-A",
        "priority_note": "【最優先】Claude がハンドメイド (再診/紹介状シナリオの gold reference)。スキーマ + PCI 後二次予防の記載粒度をご確認ください。",
    },
    {
        "id": "JC-AMI-B",
        "priority_note": "高齢女性の atypical 心筋梗塞 (心窩部痛 + 嘔気)。安全性が極めて重要 — 「胃が重い」を ACS と捉え緊急搬送判断ができるかが評価軸。",
    },
    {
        "id": "JC-CA-B",
        "priority_note": "borderline 案件。red flag (両側手根管症候群既往 + 不可逆血圧低下 + HFpEF 様症状) からの心アミロイドーシス疑いまでつなげるかが評価軸。専門性高く、subagent 生成の精度確認が必要。",
    },
    {
        "id": "JC-AHF-B",
        "priority_note": "borderline 案件。配偶者死別後 1 ヶ月のストレス + 急性発症呼吸困難。たこつぼ症候群 vs 肺炎契機 AHF vs ACS の鑑別判断、緊急入院適応の妥当性をご確認ください。",
    },
    {
        "id": "JC-AS-B",
        "priority_note": "atypical 案件。84 歳女性 frailty 高度・認知症併存、TAVI 適応 vs 保存的方針 の SDM 判断。高齢者末期意思決定の臨床妥当性をご確認ください。",
    },
    {
        "id": "JC-VSA-B",
        "priority_note": "borderline 案件。冠攣縮性狭心症の確定診断前段階での薬物選択 (β遮断薬を避けるべき) という非自明な判断が含まれているかが評価軸。",
    },
    {
        "id": "JC-SCP-B",
        "priority_note": "negative control 案件。doctor shopping の中年男性、最低限の cardiac screen のみで負荷検査を出さず心療内科紹介する判断の妥当性をご確認ください。循環器ハルシネーション (架空疾患診断) を出していないかも要確認。",
    },
]


def load_cases() -> dict:
    cases = {}
    for f in CASES_DIR.glob("JC-*.json"):
        with open(f, encoding="utf-8") as fp:
            c = json.load(fp)
        cases[c["encounter_id"]] = c
    return cases


def case_to_text_blocks(c: dict) -> list:
    """Return list of (style, text) tuples for a case."""
    blocks = []
    p = c.get("patient", {})
    e = c.get("encounter", {})
    rv = c.get("reception_vitals", {})
    kf = c.get("key_facts", {})

    # Header
    scenario_label = "シナリオ B (初診 walk-in)" if c["scenario"] == "shoshin_walkin" else "シナリオ A (再診/紹介状)"
    diff_label = "陰性対照" if c.get("is_negative_control") else f"難度: {c['difficulty']}"
    blocks.append(("h2", f"{c['encounter_id']} — {c['disease_label_jp']}"))
    blocks.append(("subtitle", f"{scenario_label} / {diff_label}"))

    # Patient
    blocks.append(("h3", "■ 患者情報"))
    pt_lines = [
        f"・年齢/性別: {p.get('age')} 歳 {p.get('gender', '')} (血液型 {p.get('blood_type', '?')})",
        f"・主訴: {e.get('chief_complaint', '')}",
    ]
    if e.get("secondary_complaints"):
        pt_lines.append(f"・副主訴: {'、'.join(e['secondary_complaints'])}")
    pt_lines += [
        f"・既往: {'、'.join(p.get('comorbidities', [])) or '—'}",
        f"・持参薬: {'、'.join(p.get('current_medications', [])) or '—'}",
        f"・アレルギー: {'、'.join(p.get('allergies', [])) or '—'}",
        f"・家族歴: {'、'.join(p.get('family_history', [])) or '—'}",
        f"・社会歴: {p.get('social_history', '—')}",
        f"・受診: {e.get('encounter_date', '')} / {e.get('department', '')} / {e.get('type', '')}",
    ]
    for line in pt_lines:
        blocks.append(("body", line))

    # Reception vitals
    if rv:
        blocks.append(("h3", "■ 受付バイタル"))
        v_parts = []
        if rv.get("BP_sys") is not None:
            v_parts.append(f"BP {rv['BP_sys']}/{rv['BP_dia']} mmHg")
        if rv.get("HR") is not None:
            v_parts.append(f"HR {rv['HR']}/min")
        if rv.get("SpO2") is not None:
            v_parts.append(f"SpO2 {rv['SpO2']}%")
        if rv.get("RR") is not None:
            v_parts.append(f"RR {rv['RR']}")
        if rv.get("BT") is not None:
            v_parts.append(f"BT {rv['BT']}℃")
        blocks.append(("body", "・" + " / ".join(v_parts)))

    # Input
    blocks.append(("h3", "■ 入力 (SLM への問診)"))
    if c["scenario"] == "shoshin_walkin":
        blocks.append(("body_label", "[口語対話 Pattern B]"))
        blocks.append(("body_quoted", c.get("input_pattern_B", "")))
    else:
        ipa = c.get("input_pattern_A", {}) or {}
        ps = ipa.get("physician_summary", {}) or {}
        for label, key in [
            ("医師サマリ (現病歴・既往整理)", "raw_text"),
            ("お薬手帳", "medication_list"),
            ("診察所見", "exam_findings"),
            ("検査結果", "lab_results"),
        ]:
            v = ps.get(key)
            if v:
                blocks.append(("body_label", f"[{label}]"))
                blocks.append(("body_quoted", v))

    # Reference SOAP
    blocks.append(("h3", "■ 正解 (reference) SOAP"))
    soap = c.get("reference_soap", {})
    for sec_label, sec in [("S (Subjective)", "S"), ("O (Objective)", "O"),
                            ("A (Assessment)", "A"), ("P (Plan)", "P")]:
        if soap.get(sec):
            blocks.append(("body_label", f"[{sec_label}]"))
            blocks.append(("body_quoted", soap[sec]))

    # Admission summary
    if c.get("reference_admission_summary"):
        blocks.append(("h3", "■ 入院時サマリ"))
        blocks.append(("body_quoted", c["reference_admission_summary"]))

    # Key facts
    blocks.append(("h3", "■ key_facts (eval runner が機械採点する正解集合)"))
    kf_lines = []
    if kf.get("expected_triage"):
        kf_lines.append(f"・期待 triage 判断: {kf['expected_triage']}")
    if kf.get("alternative_acceptable_triage"):
        kf_lines.append(f"・代替 triage: {'、'.join(kf['alternative_acceptable_triage'])}")
    if kf.get("must_not_miss"):
        kf_lines.append(f"・Must-not-miss: {'、'.join(kf['must_not_miss'])}")
    if kf.get("differential_diagnoses"):
        kf_lines.append(f"・鑑別診断: {'、'.join(kf['differential_diagnoses'])}")
    if kf.get("appropriate_next_tests"):
        kf_lines.append(f"・推奨検査: {'、'.join(kf['appropriate_next_tests'])}")
    if kf.get("diagnoses"):
        kf_lines.append(f"・確定診断: {'、'.join(kf['diagnoses'])}")
    if kf.get("diagnoses_provisional"):
        kf_lines.append(f"・暫定診断: {'、'.join(kf['diagnoses_provisional'])}")
    if kf.get("medications_to_start"):
        kf_lines.append(f"・開始薬: {'、'.join(kf['medications_to_start'])}")
    if kf.get("medications_to_continue"):
        kf_lines.append(f"・継続薬: {'、'.join(kf['medications_to_continue'])}")
    if kf.get("medications_to_stop"):
        kf_lines.append(f"・中止薬: {'、'.join(kf['medications_to_stop'])}")
    if kf.get("medications_to_consider"):
        kf_lines.append(f"・導入検討: {'、'.join(kf['medications_to_consider'])}")
    if kf.get("disposition"):
        kf_lines.append(f"・転帰: {kf['disposition']}")
    for line in kf_lines:
        blocks.append(("body", line))

    # Reviewer notes
    if c.get("notes_for_reviewer"):
        blocks.append(("h3", "■ レビュー時の評価ポイント (Claude メモ)"))
        blocks.append(("body_quoted", c["notes_for_reviewer"]))

    return blocks


def write_docx(cases: dict, selection: list):
    doc = Document()

    # Set font
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10.5)

    # Title
    title = doc.add_heading("ACI-JP-Cardio 外来カルテ生成ベンチマーク — 専門医レビュー依頼", level=0)

    # Intro
    doc.add_paragraph(
        "本書は、日本語循環器特化の小規模言語モデル (SLM) を評価するために構築した "
        "「外来カルテ生成ベンチマーク」の一部を、専門医の先生に医学的妥当性の観点から "
        "ご確認いただくための資料です。"
    )
    doc.add_paragraph(
        "ベンチマーク全体は 11 疾患 × 2 シナリオ (再診/紹介状 と 初診 walk-in) = 22 例で "
        "構成されており、本書ではその中から優先度の高い 8 例を抜粋しました。"
    )

    # Review viewpoints
    doc.add_heading("レビュー観点", level=1)
    doc.add_paragraph(
        "各症例について、以下の観点でご確認ください。問題なければ「OK」、修正が必要なら "
        "具体的にご指摘いただけますと、こちらで JSON ファイルを修正いたします。"
    )

    viewpoints = [
        ("1. 患者プロファイルの臨床的整合性",
         "年齢・性別・既往・服薬・社会歴 が疾患設定と整合しているか。不自然な組み合わせ (例: 30 代女性に重症 AS) はないか。"),
        ("2. 主訴・現病歴の臨床的妥当性",
         "主訴と疾患が矛盾していないか。受付バイタル (BP/HR/SpO2) と症状の重症度が整合しているか。"),
        ("3. 入力 (口語対話 or 医師サマリ) のリアリティ",
         "実臨床の問診らしさがあるか。AI 問診的な定型さが残っていないか。フィラー (えーと/あの) や患者の自己訂正、家族の代弁などがリアルに含まれているか。"),
        ("4. Reference SOAP の医学的妥当性",
         "・S: 入力情報を網羅し、過不足なく整理されているか。\n"
         "・O: シナリオ B (初診) では「受付バイタルのみ + 診察前段階で X なし」と明記されているか。シナリオ A では検査値・身体所見が現実的に書かれているか。\n"
         "・A: 鑑別リストが妥当か。シナリオ B では「○○疑い、>○○鑑別必要」の暫定診断スタイルになっているか。must-not-miss に重要疾患 (ACS / 解離 / PE) が含まれているか。\n"
         "・P: 検査オーダーが過不足ないか (不要な心エコー先行など、過剰検査がないか)。triage 判断 (緊急搬送 / 入院 / 外来精査 / 帰宅) が妥当か。Red flag 症状の説明が含まれているか。"),
        ("5. 薬剤選択の臨床的妥当性",
         "禁忌寄りの薬剤が処方されていないか (例: VSA でβ遮断薬、CA でβ遮断薬・ACE 阻害薬の高用量、AS で過剰利尿)。一般名と用量が現実的か。"),
        ("6. key_facts (機械採点用正解集合) の妥当性",
         "expected_triage / must-not-miss / 鑑別診断 / 推奨検査 が SOAP と整合しているか。これは eval runner が SLM 出力を採点する正解として使うので、ここがズレているとベンチマーク自体が信頼できなくなります。"),
        ("7. ハルシネーション・不適切記載",
         "架空の薬名・架空の検査値・存在しない疾患名が含まれていないか。日本の保険診療で実施できない検査・治療が記載されていないか。"),
    ]
    for h, body in viewpoints:
        p = doc.add_paragraph()
        p.add_run(h + "\n").bold = True
        p.add_run(body)

    # How to provide feedback
    doc.add_heading("修正コメントの書き方", level=1)
    doc.add_paragraph(
        "各症例の末尾に余白がございます。修正が必要な箇所を例えば次のように記載いただけますと、"
        "私が JSON を機械的に反映できます:"
    )
    examples = [
        '・「reference_soap.A の「Killip 分類 II」→「Killip I が妥当」",',
        '・「key_facts.medications_to_start に「ヘパリン」を追加してください」",',
        '・「reference_soap.P の「ニトロプルシド」は当院採用なし、「ニカルジピン静注」に変更」",',
    ]
    for ex in examples:
        doc.add_paragraph(ex, style="List Bullet")
    doc.add_paragraph(
        "完全な書き換えが必要なケースでも、要点を箇条書きで指摘いただければ私が反映します。"
    )

    # Selection summary
    doc.add_heading("レビュー対象 8 件", level=1)
    summary_para = doc.add_paragraph()
    for i, sel in enumerate(selection, 1):
        c = cases.get(sel["id"])
        if not c:
            continue
        summary_para.add_run(f"{i}. {sel['id']} — {c['disease_label_jp']}\n").bold = True

    # Per-case content
    for i, sel in enumerate(selection, 1):
        c = cases.get(sel["id"])
        if not c:
            continue
        doc.add_page_break()

        h = doc.add_heading(f"症例 {i} / 8: {sel['id']}", level=1)

        # Priority note
        p = doc.add_paragraph()
        run = p.add_run(sel["priority_note"])
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        # Case content
        for style, text in case_to_text_blocks(c):
            if style == "h2":
                doc.add_heading(text, level=2)
            elif style == "h3":
                doc.add_heading(text, level=3)
            elif style == "subtitle":
                p = doc.add_paragraph()
                p.add_run(text).italic = True
            elif style == "body":
                doc.add_paragraph(text)
            elif style == "body_label":
                p = doc.add_paragraph()
                p.add_run(text).bold = True
            elif style == "body_quoted":
                p = doc.add_paragraph(text)
                p.paragraph_format.left_indent = Cm(0.5)

        # Comment area
        doc.add_heading("先生コメント欄", level=3)
        doc.add_paragraph("(ここにレビューコメントをご記載ください)")
        doc.add_paragraph(" ")
        doc.add_paragraph(" ")
        doc.add_paragraph(" ")

    doc.save(OUT_DOCX)
    print(f"wrote {OUT_DOCX} ({OUT_DOCX.stat().st_size / 1024:.1f} KB)")


def write_txt(cases: dict, selection: list):
    """Plain text fallback for when the user wants raw copy-paste."""
    lines = []
    sep = "=" * 70
    sep_l = "-" * 70

    lines.append(sep)
    lines.append("ACI-JP-Cardio 外来カルテ生成ベンチマーク — 専門医レビュー依頼")
    lines.append(sep)
    lines.append("")
    lines.append("本書は、日本語循環器特化の小規模言語モデル (SLM) を評価するために")
    lines.append("構築した「外来カルテ生成ベンチマーク」の一部を、専門医の先生に医学的")
    lines.append("妥当性の観点からご確認いただくための資料です。")
    lines.append("")
    lines.append("ベンチマーク全体は 11 疾患 × 2 シナリオ (再診/紹介状 と 初診 walk-in)")
    lines.append("= 22 例で構成されており、本書ではその中から優先度の高い 8 例を")
    lines.append("抜粋しました。")
    lines.append("")
    lines.append(sep_l)
    lines.append("レビュー観点")
    lines.append(sep_l)
    lines.append("")
    lines.append("各症例について、以下の観点でご確認ください。問題なければ「OK」、修正が")
    lines.append("必要なら具体的にご指摘いただけますと、こちらで JSON ファイルを修正")
    lines.append("いたします。")
    lines.append("")

    for h, body in [
        ("1. 患者プロファイルの臨床的整合性",
         "年齢・性別・既往・服薬・社会歴 が疾患設定と整合しているか。不自然な組み合わせ (例: 30 代女性に重症 AS) はないか。"),
        ("2. 主訴・現病歴の臨床的妥当性",
         "主訴と疾患が矛盾していないか。受付バイタル (BP/HR/SpO2) と症状の重症度が整合しているか。"),
        ("3. 入力 (口語対話 or 医師サマリ) のリアリティ",
         "実臨床の問診らしさがあるか。AI 問診的な定型さが残っていないか。フィラーや家族の代弁がリアルか。"),
        ("4. Reference SOAP の医学的妥当性",
         "S: 過不足ないか / O: 初診なら「診察前段階」明記 / A: 鑑別リストが妥当 / P: 検査オーダーと triage 判断が妥当 / Red flag 説明あるか"),
        ("5. 薬剤選択の臨床的妥当性",
         "禁忌寄りの薬剤が処方されていないか (VSA でβ遮断薬、CA でβ遮断薬高用量、AS で過剰利尿等)。一般名と用量が現実的か。"),
        ("6. key_facts (機械採点用正解集合) の妥当性",
         "expected_triage / must-not-miss / 鑑別 / 推奨検査 が SOAP と整合しているか。"),
        ("7. ハルシネーション・不適切記載",
         "架空の薬名・架空の検査値・存在しない疾患名が含まれていないか。"),
    ]:
        lines.append(f"【{h}】")
        lines.append(body)
        lines.append("")

    lines.append(sep_l)
    lines.append("修正コメントの書き方")
    lines.append(sep_l)
    lines.append("")
    lines.append("各症例の末尾に「先生コメント欄」がございます。修正が必要な箇所を")
    lines.append("例えば次のように記載いただけますと反映できます:")
    lines.append("")
    lines.append("  ・reference_soap.A の「Killip 分類 II」→「Killip I が妥当」")
    lines.append("  ・key_facts.medications_to_start に「ヘパリン」を追加")
    lines.append("  ・reference_soap.P の「ニトロプルシド」を「ニカルジピン静注」に変更")
    lines.append("")

    lines.append(sep_l)
    lines.append("レビュー対象 8 件")
    lines.append(sep_l)
    for i, sel in enumerate(selection, 1):
        c = cases.get(sel["id"])
        if not c:
            continue
        lines.append(f"  {i}. {sel['id']} — {c['disease_label_jp']}")
    lines.append("")

    for i, sel in enumerate(selection, 1):
        c = cases.get(sel["id"])
        if not c:
            continue
        lines.append("")
        lines.append(sep)
        lines.append(f"症例 {i} / 8: {sel['id']}")
        lines.append(sep)
        lines.append("")
        lines.append(f"<<<優先度メモ>>>")
        lines.append(sel["priority_note"])
        lines.append("")

        for style, text in case_to_text_blocks(c):
            if style == "h2":
                lines.append("")
                lines.append(text)
            elif style == "h3":
                lines.append("")
                lines.append(text)
            elif style == "subtitle":
                lines.append(f"({text})")
            elif style == "body":
                lines.append(text)
            elif style == "body_label":
                lines.append("")
                lines.append(text)
            elif style == "body_quoted":
                # Indent each line of the body
                for ln in text.split("\n"):
                    lines.append("  " + ln)

        lines.append("")
        lines.append("【先生コメント欄】")
        lines.append("(ここにレビューコメントをご記載ください)")
        lines.append("")
        lines.append("")
        lines.append("")
        lines.append("")

    OUT_TXT.write_text("\n".join(lines), encoding="utf-8")
    print(f"wrote {OUT_TXT} ({OUT_TXT.stat().st_size / 1024:.1f} KB)")


def main():
    cases = load_cases()
    write_docx(cases, SELECTION)
    write_txt(cases, SELECTION)


if __name__ == "__main__":
    main()
